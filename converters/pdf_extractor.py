import os
import zipfile
from PyPDF2 import PdfReader

class PDFExtractor:
    def __init__(self):
        self.output_dir = "temp/output"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def extract_images(self, pdf_path: str, session_id: str) -> str:
        """Extract images from PDF (placeholder with info)"""
        output_path = os.path.join(self.output_dir, f"{session_id}_images.zip")
        
        # Create informative zip file
        with zipfile.ZipFile(output_path, 'w') as zipf:
            info_content = f"""SaveMedia PDF Converter - Image Extraction

PDF File: {os.path.basename(pdf_path)}
Extraction ID: {session_id}

This ZIP file contains information about image extraction.

For actual image extraction from PDF files, please use:
- Our desktop application
- Advanced online tools
- Professional PDF software

Thank you for using SaveMedia PDF Converter!

Features available:
✓ PDF to Text extraction
✓ PDF to DOCX conversion
✓ Multiple file combination
✓ Secure file processing

Visit our website for more tools and features.
"""
            zipf.writestr("README.txt", info_content)
            zipf.writestr("image_extraction_info.txt", 
                         "Advanced image extraction requires additional libraries.\n"
                         "Basic text extraction is available in the Text format option.")
        
        return output_path
    
    def extract_text(self, pdf_path: str, session_id: str) -> str:
        """Extract text from PDF"""
        output_path = os.path.join(self.output_dir, f"{session_id}.txt")
        
        try:
            with open(pdf_path, 'rb') as file:
                reader = PdfReader(file)
                
                text_content = f"""SaveMedia PDF Converter - Text Extraction
===========================================

PDF File: {os.path.basename(pdf_path)}
Total Pages: {len(reader.pages)}
Extraction ID: {session_id}

===========================================
EXTRACTED TEXT CONTENT:
===========================================

"""
                
                for page_num, page in enumerate(reader.pages, 1):
                    text = page.extract_text()
                    if text.strip():
                        text_content += f"\n{'='*50}\n"
                        text_content += f"PAGE {page_num}\n"
                        text_content += f"{'='*50}\n\n"
                        text_content += text + "\n"
                
                text_content += f"\n\n{'='*50}\n"
                text_content += "END OF DOCUMENT\n"
                text_content += "Converted by SaveMedia PDF Converter\n"
                text_content += "Thank you for using our service!"
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                
                return output_path
                
        except Exception as e:
            raise Exception(f"Text extraction failed: {str(e)}")
    
    def convert_to_docx(self, pdf_path: str, session_id: str) -> str:
        """Convert PDF to DOCX format"""
        output_path = os.path.join(self.output_dir, f"{session_id}.docx")
        
        try:
            # First extract text
            text_path = self.extract_text(pdf_path, session_id + "_temp")
            
            try:
                # Try to create actual DOCX using python-docx
                from docx import Document
                from docx.shared import Inches
                
                doc = Document()
                
                # Add title
                title = doc.add_heading('PDF to DOCX Conversion', 0)
                
                # Add file info
                doc.add_paragraph(f"Original PDF: {os.path.basename(pdf_path)}")
                doc.add_paragraph(f"Conversion ID: {session_id}")
                doc.add_paragraph("Converted by SaveMedia PDF Converter")
                doc.add_paragraph("")
                
                # Read extracted text
                with open(text_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Add content
                for line in content.split('\n'):
                    if line.strip() and not line.startswith('='):
                        doc.add_paragraph(line.strip())
                
                doc.save(output_path)
                return output_path
                
            except ImportError:
                # If python-docx is not available, rename text file to .docx
                os.rename(text_path, output_path)
                return output_path
            finally:
                # Cleanup temp file if it exists and is different from output
                if os.path.exists(text_path) and text_path != output_path:
                    os.remove(text_path)
                    
        except Exception as e:
            raise Exception(f"DOCX conversion failed: {str(e)}")
